#!/usr/bin/env python3
"""Build a professional Data Dictionary document (.docx) from a constrained
Markdown file.

Usage:
    python build_docx.py <input.md> <output.docx>

Requires:  pip install python-docx

Supported Markdown subset (anything else is rendered as plain text):

  ---  frontmatter  ---   title / subtitle / project / version / date / author /
                          status  -> rendered as a cover page + document header
  # .. ####               headings 1-4 (every "#" heading starts a new page)
  plain paragraphs        with **bold**, *italic* and `code` inline spans
  - item / * item         bullet lists (2 leading spaces per nesting level)
  1. item                 numbered lists
  | a | b |               pipe tables; first row is the header row
  ![caption](path)        centered image with an auto-numbered "Figure N" caption
  ```lang ... ```         code block (monospace, light-grey shading)
  > text                  callout / quote block

Exit code is non-zero when a referenced image file is missing, so the caller
can fix the path and re-run.
"""

import os
import re
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Emu, Inches, Pt, RGBColor
except ImportError:
    sys.exit("python-docx is not installed. Run:  pip install python-docx")

ACCENT = RGBColor(0x1F, 0x38, 0x64)      # dark blue used for headings
ACCENT_HEX = "1F3864"
TABLE_HEADER_HEX = "2E5496"
CODE_SHADE_HEX = "F2F2F2"
GREY = RGBColor(0x59, 0x59, 0x59)
MAX_IMG_WIDTH = Inches(6.2)

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
IMAGE_RE = re.compile(r"^!\[(?P<caption>.*?)\]\((?P<path>[^)]+)\)\s*$")
BULLET_RE = re.compile(r"^(?P<indent>\s*)[-*]\s+(?P<text>.+)$")
NUMBER_RE = re.compile(r"^(?P<indent>\s*)\d+[.)]\s+(?P<text>.+)$")
HEADING_RE = re.compile(r"^(?P<hashes>#{1,4})\s+(?P<text>.+)$")


# ---------------------------------------------------------------- low level

def _field(paragraph, instr, placeholder=None):
    """Append a Word field (e.g. PAGE, TOC) to a paragraph."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    run._r.append(begin)
    run._r.append(instr_el)
    if placeholder is not None:
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        run._r.append(sep)
        ph_run = paragraph.add_run(placeholder)
        ph_run.font.italic = True
        ph_run.font.color.rgb = GREY
        run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _shade_cell(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _shade_paragraph(paragraph, hex_fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    p_pr.append(shd)


def _mark_update_fields(doc):
    """Make Word offer to refresh the TOC when the document is opened."""
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    doc.settings.element.append(upd)


def add_inline(paragraph, text, bold=False, color=None, size=None):
    """Write text with **bold**, *italic* and `code` spans into a paragraph."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        run = None
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token)
        if bold:
            run.bold = True
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = size


# ---------------------------------------------------------------- styling

def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    sizes = {1: 16, 2: 13, 3: 12, 4: 11}
    for level, size in sizes.items():
        st = doc.styles[f"Heading {level}"]
        st.font.color.rgb = ACCENT
        st.font.size = Pt(size)
        st.font.bold = True
        if level == 4:
            st.font.italic = True
    doc.styles["Heading 1"].paragraph_format.page_break_before = True
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(0)


def build_cover(doc, meta):
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta.get("title", "Data Dictionary"))
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta.get("subtitle", "Data Dictionary"))
    run.font.size = Pt(16)
    run.font.color.rgb = GREY
    doc.add_paragraph()
    doc.add_paragraph()

    rows = [(k.capitalize(), meta[k]) for k in
            ("project", "version", "date", "author", "status") if meta.get(k)]
    if rows:
        table = doc.add_table(rows=len(rows), cols=2, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for i, (label, value) in enumerate(rows):
            left, right = table.rows[i].cells
            left.width, right.width = Inches(1.6), Inches(3.2)
            p = left.paragraphs[0]
            r = p.add_run(label)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(left, TABLE_HEADER_HEX)
            right.paragraphs[0].add_run(str(value))
    doc.add_page_break()


def build_toc(doc):
    head = doc.add_paragraph()
    run = head.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT
    toc = doc.add_paragraph()
    _field(toc, r'TOC \o "1-3" \h \z \u',
           placeholder="Table of contents — in Word press Ctrl+A then F9 to populate.")
    doc.add_page_break()


def build_header_footer(doc, meta):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run(meta.get("title", "Data Dictionary"))
    run.font.size = Pt(9)
    run.font.color.rgb = GREY

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    _field(footer_p, "PAGE")
    run = footer_p.add_run(" of ")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    _field(footer_p, "NUMPAGES")
    for r in footer_p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = GREY


# ---------------------------------------------------------------- markdown

def parse_frontmatter(lines):
    meta = {}
    if not lines or lines[0].strip() != "---":
        return meta, lines
    for i in range(1, len(lines)):
        if lines[i].strip() == "---":
            for raw in lines[1:i]:
                if ":" in raw:
                    key, _, value = raw.partition(":")
                    meta[key.strip().lower()] = value.strip().strip('"')
            return meta, lines[i + 1:]
    return meta, lines


def add_image(doc, path, caption, base_dir, figure_no, warnings):
    full = path if os.path.isabs(path) else os.path.join(base_dir, path)
    full = os.path.normpath(full)
    if not os.path.isfile(full):
        warnings.append(f"missing image: {path}")
        p = doc.add_paragraph()
        run = p.add_run(f"[Image not found: {path}]")
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        return
    doc.add_picture(full)
    shape = doc.inline_shapes[-1]
    if shape.width > MAX_IMG_WIDTH:
        ratio = MAX_IMG_WIDTH / shape.width
        shape.height = Emu(int(shape.height * ratio))
        shape.width = MAX_IMG_WIDTH
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Figure {figure_no}: {caption}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GREY


def add_table(doc, rows):
    parsed = []
    for raw in rows:
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue  # separator row
        parsed.append(cells)
    if not parsed:
        return
    n_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=n_cols, style="Table Grid")
    table.autofit = True
    for i, cells in enumerate(parsed):
        for j in range(n_cols):
            text = cells[j] if j < len(cells) else ""
            cell = table.rows[i].cells[j]
            p = cell.paragraphs[0]
            if i == 0:
                _shade_cell(cell, TABLE_HEADER_HEX)
                add_inline(p, text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                           size=Pt(10))
            else:
                add_inline(p, text, size=Pt(10))
            for r in p.runs:
                r.font.size = Pt(10)
    doc.add_paragraph()


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    _shade_paragraph(p, CODE_SHADE_HEX)
    p.paragraph_format.left_indent = Inches(0.15)
    for i, line in enumerate(code_lines):
        if i:
            p.runs[-1].add_break()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def render_body(doc, lines, base_dir, warnings):
    figure_no = 0
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):                       # code block
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i].rstrip("\n"))
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        if stripped.startswith("|"):                         # table
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            add_table(doc, rows)
            continue

        m = HEADING_RE.match(stripped)
        if m:
            doc.add_heading(m.group("text").strip(), level=len(m.group("hashes")))
            i += 1
            continue

        m = IMAGE_RE.match(stripped)
        if m:
            figure_no += 1
            add_image(doc, m.group("path").strip(), m.group("caption").strip(),
                      base_dir, figure_no, warnings)
            i += 1
            continue

        m = BULLET_RE.match(line)
        if m:
            level = min(len(m.group("indent")) // 2, 2)
            style = "List Bullet" if level == 0 else f"List Bullet {level + 1}"
            p = doc.add_paragraph(style=style)
            add_inline(p, m.group("text"))
            i += 1
            continue

        m = NUMBER_RE.match(line)
        if m:
            level = min(len(m.group("indent")) // 2, 2)
            style = "List Number" if level == 0 else f"List Number {level + 1}"
            p = doc.add_paragraph(style=style)
            add_inline(p, m.group("text"))
            i += 1
            continue

        if stripped.startswith(">"):                         # callout
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        if stripped in ("", "---"):
            i += 1
            continue

        p = doc.add_paragraph()                              # plain paragraph
        add_inline(p, stripped)
        i += 1


# ---------------------------------------------------------------- main

def main():
    if len(sys.argv) != 3:
        sys.exit("Usage: python build_docx.py <input.md> <output.docx>")
    md_path, out_path = sys.argv[1], sys.argv[2]
    if not os.path.isfile(md_path):
        sys.exit(f"Input file not found: {md_path}")

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()
    meta, body = parse_frontmatter(lines)
    base_dir = os.path.dirname(os.path.abspath(md_path))

    doc = Document()
    style_document(doc)
    build_header_footer(doc, meta)
    build_cover(doc, meta)
    build_toc(doc)
    warnings = []
    render_body(doc, body, base_dir, warnings)
    _mark_update_fields(doc)

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")

    if warnings:
        print("\nWARNINGS — fix these and re-run:")
        for w in warnings:
            print(f"  - {w}")
        sys.exit(1)


if __name__ == "__main__":
    main()
